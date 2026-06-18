"""
Document export.

Turns the structured report (the dict from `render.render_dict`) into the kind of
artefact a reviewer actually files: a PDF and an editable DOCX. Both consume the same
payload, so they stay consistent with the Markdown/JSON outputs — including the
References section that resolves each citation to its source.

    to_pdf(report, "report.pdf")    # via reportlab  (pure-python, no system libs)
    to_docx(report, "report.docx")  # via python-docx (editable)

Returns the file as bytes too (UI download buttons), so nothing has to touch disk.
"""

import io


def _title(report: dict) -> str:
    return f"Insurance Risk Report — {report['meta'].get('site_name', '')}"


def _subtitle(report: dict) -> str:
    m = report["meta"]
    return (f"Inspected {m.get('inspection_date', '')} by {m.get('inspector', '')} · "
            f"jurisdiction {m.get('jurisdiction', '')}")


def to_pdf(report: dict, path: str | None = None) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    except ImportError as e:
        raise RuntimeError("PDF export needs reportlab: pip install reportlab.") from e

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=18 * mm, bottomMargin=18 * mm)
    styles = getSampleStyleSheet()
    h2 = ParagraphStyle("h2b", parent=styles["Heading2"], spaceBefore=10)
    flow = [Paragraph(_title(report), styles["Title"]),
            Paragraph(_subtitle(report), styles["Italic"]), Spacer(1, 8)]

    for field in report["fields"]:
        flow.append(Paragraph(field["name"].replace("_", " ").title(), h2))
        for claim in field["claims"]:
            cite = f' <font color="#666">[{", ".join(claim["citations"])}]</font>' if claim["citations"] else ""
            flow.append(Paragraph(f"• {claim['text']}{cite}", styles["BodyText"]))

    s = report["verification"]["summary"]
    flow.append(Paragraph("Verification Summary", h2))
    flow.append(Paragraph(
        f"Supported: {s['SUPPORTED']} | Partial: {s['PARTIAL']} | "
        f"Unsupported: {s['UNSUPPORTED']} | No citation: {s['NO_CITATION']}", styles["BodyText"]))
    flags = [c for c in report["verification"]["checks"]
             if c["status"] in ("UNSUPPORTED", "NO_CITATION")]
    if flags:
        flow.append(Paragraph("Flagged for human review:", styles["BodyText"]))
        for c in flags:
            flow.append(Paragraph(f"• [{c['status']}] ({c['field']}) {c['claim']} — {c['reason']}",
                                  styles["BodyText"]))

    if report.get("references"):
        flow.append(Paragraph("References", h2))
        for r in report["references"]:
            tail = f" — {r['source_ref']}" if r["source_ref"] else ""
            flow.append(Paragraph(f"[{r['id']}] {r['title']}{tail}", styles["BodyText"]))

    doc.build(flow)
    data = buf.getvalue()
    if path:
        with open(path, "wb") as f:
            f.write(data)
    return data


def to_docx(report: dict, path: str | None = None) -> bytes:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("DOCX export needs python-docx: pip install python-docx.") from e

    doc = Document()
    doc.add_heading(_title(report), level=0)
    doc.add_paragraph(_subtitle(report)).italic = True

    for field in report["fields"]:
        doc.add_heading(field["name"].replace("_", " ").title(), level=1)
        for claim in field["claims"]:
            cite = f"  [{', '.join(claim['citations'])}]" if claim["citations"] else ""
            doc.add_paragraph(f"{claim['text']}{cite}", style="List Bullet")

    s = report["verification"]["summary"]
    doc.add_heading("Verification Summary", level=1)
    doc.add_paragraph(
        f"Supported: {s['SUPPORTED']} | Partial: {s['PARTIAL']} | "
        f"Unsupported: {s['UNSUPPORTED']} | No citation: {s['NO_CITATION']}")
    flags = [c for c in report["verification"]["checks"]
             if c["status"] in ("UNSUPPORTED", "NO_CITATION")]
    if flags:
        doc.add_paragraph("Flagged for human review:")
        for c in flags:
            doc.add_paragraph(f"[{c['status']}] ({c['field']}) {c['claim']} — {c['reason']}",
                              style="List Bullet")

    if report.get("references"):
        doc.add_heading("References", level=1)
        for r in report["references"]:
            tail = f" — {r['source_ref']}" if r["source_ref"] else ""
            doc.add_paragraph(f"[{r['id']}] {r['title']}{tail}", style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    data = buf.getvalue()
    if path:
        with open(path, "wb") as f:
            f.write(data)
    return data
